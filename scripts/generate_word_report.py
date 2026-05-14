#!/usr/bin/env python3
"""
Generate Word (.docx) technical report.
Output: GCC AI Intelligence Platform – Yahya Alwashahi.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

NAVY  = RGBColor(0x1B, 0x4F, 0x72)
GOLD  = RGBColor(0xC3, 0x9B, 0x4E)
GREEN = RGBColor(0x1A, 0x7A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x09, 0x13, 0x20)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '1B4F72')
        tcBorders.append(border)
    tcPr.append(tcBorders)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default paragraph style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

# ── Cover ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GCC AI Intelligence Platform')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Technical Report — Competition Submission')
r2.font.size = Pt(16)
r2.font.color.rgb = GOLD
r2.bold = True

doc.add_paragraph()
meta = [
    ('Platform', 'GCC AI Strategic Intelligence Platform'),
    ('Focus', 'AI-Powered Youth Employment Forecasting & Policy Decision Support for the Gulf Cooperation Council'),
    ('Participant', 'Yahya Alwashahi'),
    ('Data Source', 'World Bank Open Data API v2'),
    ('Coverage', '6 GCC Nations · 5 Indicators · 2010–2024'),
    ('Languages', 'English · Arabic (Ministry-Grade Bilingual Output)'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p.add_run(f'{label}: ')
    run_l.bold = True
    run_l.font.color.rgb = NAVY
    run_v = p.add_run(value)
    run_v.font.color.rgb = DARK

doc.add_page_break()

# ── Section helper ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    # Gold underline rule
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'C39B4E')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    # parse **bold** inline
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            p.add_run(part)
    p.paragraph_format.space_after = Pt(3)

def table_from_rows(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(hdr_cells[i], '1B4F72')
    # data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            if ri % 2 == 0:
                set_cell_bg(cells[ci], 'EEF2F7')
    doc.add_paragraph()

# ── 1. Executive Summary ──────────────────────────────────────────────────────
h1('1. Executive Summary')
body('The GCC AI Intelligence Platform is an end-to-end AI-powered statistical intelligence system designed to support youth employment forecasting and strategic policy planning across the six Gulf Cooperation Council nations: Saudi Arabia, UAE, Qatar, Kuwait, Bahrain, and Oman.')
body('The platform combines real World Bank Open Data, a six-model ensemble forecasting engine with automatic model selection, an AI-generated bilingual executive intelligence layer, elasticity-based policy scenario simulation, and a complete explainability and trust framework — all delivered through a production-grade interactive dashboard.')
h2('Key Capabilities')
for cap in [
    '**Real-time GCC regional intelligence dashboard** with AI risk classification',
    '**Multi-model time-series forecasting** with expanding-window cross-validation',
    '**Bilingual (English/Arabic) AI-generated ministerial intelligence reports**',
    '**Eight curated strategic policy scenario presets** with GCC-wide impact analysis',
    '**Transparent AI:** trust scoring, driver intelligence, confidence classification, and a full Responsible AI methodology centre',
    '**Fully offline-capable** with pre-seeded World Bank data cache',
]:
    bullet(cap)

# ── 2. Problem Statement ──────────────────────────────────────────────────────
h1('2. Problem Statement / Challenge')
h2('2.1 The GCC Youth Employment Challenge')
body('Youth unemployment is one of the most structurally significant economic challenges facing the Gulf Cooperation Council. Across the six member states, youth unemployment rates range from under 5% (UAE, Qatar) to over 25% (Saudi Arabia, Oman), with significant variation driven by differing economic structures, nationalisation policy trajectories, educational investment levels, and demographic pressures.')
h2('Key Structural Drivers')
for d in [
    'Persistent skills mismatches between university graduates and private-sector demand',
    'Dependency on public-sector employment as the default absorber of national youth labour',
    'Rapid demographic growth increasing the youth cohort size faster than job creation',
    'Digital economy transitions requiring reskilling pipelines that remain immature',
    'Regional oil-price cycles creating fiscal volatility that affects public-sector hiring capacity',
]:
    bullet(d)

h2('2.2 The Policy Intelligence Gap')
body('Despite the urgency, GCC policymakers face a persistent intelligence gap across five dimensions:')
gaps = [
    ('Fragmented data', 'Labour market statistics are scattered across national statistical agencies, the World Bank, the ILO, and OECD databases with inconsistent publication timelines'),
    ('No regional comparative view', 'Governments lack a single platform enabling real-time GCC peer benchmarking'),
    ('No forward-looking intelligence', 'Current reporting is retrospective; strategic planning lacks AI-powered predictive signals'),
    ('No scenario modelling', 'Policymakers cannot model the employment impact of specific interventions'),
    ('No bilingual executive layer', 'Insights generated in English are inaccessible to Arabic-speaking ministry audiences without manual translation'),
]
for label, desc in gaps:
    bullet(f'**{label}** — {desc}')
body('This platform was built to close all five gaps simultaneously.')

# ── 3. Project Objectives ─────────────────────────────────────────────────────
h1('3. Project Objectives')
table_from_rows(
    ['Objective', 'Outcome'],
    [
        ('Aggregate real GCC labour-market data', 'Live World Bank API pipeline with offline-capable cache'),
        ('Enable regional comparative intelligence', 'Six-country dashboard with rankings, heatmaps, and GCC averages'),
        ('Deliver AI-powered forward-looking forecasts', 'Six-model ensemble with automatic selection and prediction intervals'),
        ('Generate interpretable, actionable intelligence', 'AI-written bilingual executive reports with risk classification'),
        ('Enable policy scenario modelling', 'Eight strategic presets with elasticity-based impact simulation'),
        ('Build institutional trust through transparency', 'Trust scoring, confidence classification, driver narratives, Responsible AI centre'),
        ('Support Arabic-first ministry reporting', 'RTL-formatted Arabic executive reports downloadable in standard formats'),
    ]
)

# ── 4. Project Idea & Creative Aspect ─────────────────────────────────────────
h1('4. Project Idea & Creative Aspect')
body('The GCC AI Intelligence Platform reframes the challenge of GCC youth employment not as a data problem, but as an intelligence gap. The creative insight is that policymakers do not need more charts — they need a system that transforms data into decisions.')
h2('Creative Differentiators')
for d in [
    '**Ministry-grade bilingual AI reports** — Arabic executive intelligence generated in parallel with English, not as translation, but as a first-class output designed for Arabic-speaking government audiences',
    '**AI that explains itself** — A complete trust and explainability layer with a 0-100 trust score, five confidence tiers, and driver narratives that explain why a forecast is reliable or not',
    '**Policy scenario simulator** — Policymakers can test the impact of interventions (e.g., +5% digital investment) across all 6 GCC nations simultaneously, seeing which countries benefit most',
    '**Risk classification system** — Eight dynamic risk labels (Stable, Moderate Risk, High Risk, Recovery Phase, etc.) assigned automatically per country/indicator, enabling instant strategic triage',
    '**Demo-safe, deployment-ready** — The platform runs fully offline with pre-seeded real data; judges see a live production system, not a prototype',
]:
    bullet(d)

# ── 5. Methodology ─────────────────────────────────────────────────────────────
h1('5. Methodology / Implementation Mechanism')
h2('5.1 Data Pipeline')
body('Real World Bank Open Data is fetched via the API v2 endpoint (https://api.worldbank.org/v2/) for all 6 GCC nations across 5 indicators from 2010–2024. Data is cleaned (linear interpolation for isolated missing values), validated, and cached to CSV files for offline operation.')

h2('5.2 Forecasting Engine')
body('Six candidate models are evaluated for every forecast request using expanding-window cross-validation:')
table_from_rows(
    ['Model', 'Type', 'Strength'],
    [
        ('Naïve', 'Baseline', 'Last-value persistence'),
        ('Seasonal Naïve', 'Baseline', 'Seasonal pattern replication'),
        ('Moving Average', 'Baseline', 'Local smoothing'),
        ('Drift', 'Baseline', 'Linear trend extrapolation'),
        ('SARIMAX', 'Statistical', 'Seasonal autoregression with AIC order selection'),
        ('LightGBM', 'Machine Learning', 'Gradient boosting with lag/rolling/calendar features'),
    ]
)
body('Model selection is fully automatic — sMAPE (Symmetric Mean Absolute Percentage Error) on holdout folds determines the winner.')

h2('5.3 AI Intelligence Layer')
body('Beyond numerical forecasts, a custom AI engine (src/intelligence.py) generates structured intelligence narratives including: dynamic risk classification, strategic alert system, causal driver interpretation, and comparative GCC intelligence — all in English and Arabic.')

h2('5.4 Scenario Simulation')
body('Policy lever adjustments are applied to baseline forecasts through an elasticity model. Eight curated scenario presets cover the full range of GCC policy planning situations, generating GCC-wide comparative analyses showing which countries benefit most or least under each scenario.')

h2('5.5 Explainability & Trust Layer')
body('A composite trust score (0–100) is computed per forecast from four signals: sMAPE accuracy, forecast horizon, prediction interval width, and series volatility. Five confidence tiers (High Confidence to Unstable Forecast) are assigned automatically, with bilingual explanation provided to users.')

# ── 6. Tools & Technologies ────────────────────────────────────────────────────
h1('6. Tools & Technologies Used')
table_from_rows(
    ['Category', 'Technology', 'Purpose'],
    [
        ('Dashboard Framework', 'Streamlit', 'Interactive web application'),
        ('Data Processing', 'Pandas, NumPy', 'Data manipulation and numerical computation'),
        ('Visualisation', 'Plotly', 'Interactive charts and heatmaps'),
        ('Statistical Forecasting', 'Statsmodels (SARIMAX)', 'Seasonal autoregressive modelling'),
        ('Machine Learning', 'LightGBM', 'Gradient boosting forecasting with quantile regression'),
        ('Data Source', 'World Bank Open Data API v2', 'Real GCC labour-market statistics'),
        ('Deployment', 'Streamlit Cloud', 'Live public platform hosting'),
        ('Version Control', 'GitHub', 'Source code management'),
        ('Language', 'Python 3.11', 'Primary development language'),
        ('Document Generation', 'python-pptx, python-docx', 'Submission file generation'),
    ]
)

# ── 7. Results & Outputs ───────────────────────────────────────────────────────
h1('7. Results & Outputs')
h2('7.1 Platform Scope')
table_from_rows(
    ['Metric', 'Value'],
    [
        ('GCC Countries covered', '6 (Saudi Arabia, UAE, Qatar, Kuwait, Bahrain, Oman)'),
        ('Indicators', '5 World Bank indicators'),
        ('Data range', '2010–2024 (15 annual observations per series)'),
        ('Forecast horizon', '1–5 years (annual)'),
        ('Models evaluated per forecast', '6'),
        ('Scenario presets', '8'),
        ('Intelligence report sections per forecast', '8 (EN) + 8 (AR)'),
        ('Risk profile labels', '8'),
        ('Output languages', '2 (English + Arabic)'),
        ('Live URL', 'https://competition-efvjqtcbnxcvt3te6fqkaw.streamlit.app/'),
    ]
)

h2('7.2 Sample Intelligence Output')
body('For Saudi Arabia, Youth Unemployment Rate:')
for item in [
    '**Risk Profile:** Moderate Risk',
    '**Trend:** Improving (−0.8pp YoY)',
    '**Strategic Alerts:** 3 active (YoY Signal, GCC Divergence, Target Range breach)',
    '**Model Selected:** Moving Average, sMAPE 10.4% (Tier B — Good)',
    '**Trust Score:** 69/100 — Elevated Uncertainty',
    '**Forecast (3yr):** Projected improvement trajectory with 80% prediction interval',
]:
    bullet(item)

h2('7.3 Societal Impact')
body('The platform directly supports the stated workforce development objectives of all six GCC national visions (Vision 2030, UAE Centennial 2071, QNV 2030, Kuwait Vision 2035, Bahrain Economic Vision 2030, Oman Vision 2040).')
body('Immediate impact: Better-informed, evidence-based policy. Medium-term: Regional coordination through GCC peer benchmarking. Long-term: Building institutional AI capacity within Gulf governments.')

# ── 8. Conclusion ─────────────────────────────────────────────────────────────
h1('8. Conclusion')
body('The GCC AI Intelligence Platform represents a practical, deployable AI-assisted decision support system specifically designed for the governance realities of Gulf Cooperation Council member states. It moves beyond conventional analytics dashboards by generating intelligence — not just charts — closing the scenario modelling gap, building institutional trust through transparency, serving Arabic-first audiences, and operating fully offline.')
body('The platform is live, accessible, and ready for institutional evaluation at:')
p = doc.add_paragraph()
r = p.add_run('https://competition-efvjqtcbnxcvt3te6fqkaw.streamlit.app/')
r.bold = True
r.font.color.rgb = NAVY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GCC AI Intelligence Platform · Technical Report · Yahya Alwashahi')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
r.italic = True

out = '/home/user/competition/GCC AI Intelligence Platform – Yahya Alwashahi.docx'
doc.save(out)
print(f'Saved: {out}')
