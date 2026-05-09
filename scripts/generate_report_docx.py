#!/usr/bin/env python3
"""Generate Technical Report as DOCX and PDF from markdown content."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

NAVY   = RGBColor(0x1B, 0x4F, 0x72)
GOLD   = RGBColor(0xC3, 0x9B, 0x4E)
DARK   = RGBColor(0x09, 0x13, 0x20)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF4, 0xF6, 0xF9)
MGRAY  = RGBColor(0x77, 0x77, 0x77)
GREEN  = RGBColor(0x1A, 0x7A, 0x4A)

OUTPUT_DIR = "presentation"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_spacing(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def add_run(para, text, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    set_para_spacing(para, before=14 if level == 1 else 8, after=4)
    if level == 1:
        add_run(para, text, bold=True, size=16, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), 'C39B4E')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        add_run(para, text, bold=True, size=13, color=NAVY)
    elif level == 3:
        add_run(para, text, bold=True, size=11, color=DARK)
    return para

def add_body(doc, text, indent=False, italic=False, color=None, size=10.5):
    para = doc.add_paragraph()
    set_para_spacing(para, before=2, after=3)
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    add_run(para, text, italic=italic, size=size, color=color)
    return para

def add_bullet(doc, text, level=0, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        add_run(para, bold_prefix, bold=True, size=10.5, color=NAVY)
        add_run(para, text, size=10.5)
    else:
        add_run(para, text, size=10.5)
    return para

def add_code(doc, code_text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=4, after=4)
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF2F7')
    pPr.append(shd)
    return para

def add_table(doc, headers, rows, col_widths=None, header_bg='1B4F72'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)
    for r_idx, row in enumerate(rows):
        drow = table.rows[r_idx + 1]
        bg = 'F4F6F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = drow.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            set_cell_bg(cell, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

def build_docx(output_path):
    doc = Document()
    set_page_margins(doc)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    # Title page
    bar = doc.add_paragraph()
    set_para_spacing(bar, before=0, after=0)
    bar_run = bar.add_run('█' * 120)
    bar_run.font.color.rgb = GOLD
    bar_run.font.size = Pt(4)
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title, before=30, after=6)
    add_run(title, 'GCC AI Intelligence Platform', bold=True, size=28, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(subtitle, before=0, after=20)
    add_run(subtitle, 'Technical Report — Competition Submission', bold=False, size=14, color=GOLD)
    for label, value in [
        ('Platform', 'GCC AI Strategic Intelligence Platform'),
        ('Focus', 'AI-Powered Youth Employment Forecasting & Policy Decision Support'),
        ('Data Source', 'World Bank Open Data API v2'),
        ('Coverage', '6 GCC Nations · 5 Indicators · 2010–2024'),
        ('Languages', 'English · Arabic (Ministry-Grade Bilingual Output)'),
        ('Category', 'Competition Submission'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before=2, after=2)
        add_run(p, f'{label}:  ', bold=True, size=11, color=NAVY)
        add_run(p, value, size=11, color=DARK)
    doc.add_page_break()
    doc.save(output_path)
    print(f'✅  DOCX saved: {output_path}')
    return output_path

if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, 'GCC_AI_Intelligence_Platform_Technical_Report.docx')
    build_docx(docx_path)
